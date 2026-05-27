#!/usr/bin/env python3
from __future__ import annotations

import argparse
import base64
import json
import re
import shutil
import subprocess
import sys
from datetime import date
from html import escape
from pathlib import Path
from typing import Any

from current_state_clarity import DIMENSION_LABELS as CLARITY_DIMENSION_LABELS
from current_state_clarity import LEVEL_LABELS as CLARITY_LEVEL_LABELS
from current_state_clarity import assess_current_state_clarity

try:
    from docx import Document
    from docx.shared import Pt
except Exception:  # pragma: no cover
    Document = None
    Pt = None

try:
    from weasyprint import HTML
except Exception:  # pragma: no cover
    HTML = None


ROOT = Path(__file__).resolve().parent.parent
CSS_PATH = ROOT / "templates" / "report-theme.css"
ICEBERG_PHOTO_PATH = ROOT / "assets" / "iceberg-user-provided.png"

WEIGHTS = {
    "goal_impact": 0.25,
    "causal_leverage": 0.20,
    "stage_urgency": 0.15,
    "resource_constraint": 0.15,
    "changeability": 0.10,
    "spillover_risk": 0.10,
    "evidence_strength": 0.05,
}

DIMENSION_LABELS = {
    "goal_impact": "目标影响",
    "causal_leverage": "因果牵引",
    "stage_urgency": "阶段紧迫",
    "resource_constraint": "资源约束",
    "changeability": "可改变性",
    "spillover_risk": "风险外溢",
    "evidence_strength": "证据强度",
}

SOURCE_LABELS = {
    "observed": "已观察",
    "estimated": "估计",
    "assumed": "假设",
}

CONFIDENCE_LABELS = {
    "low": "低",
    "medium": "中",
    "medium-high": "中高",
}

NAV_ITEMS = [
    ("summary", "结论"),
    ("situation", "现状"),
    ("current-state", "清楚了吗"),
    ("visuals", "图解"),
    ("contradictions", "判断过程"),
    ("principal", "主要矛盾"),
    ("secondary", "次要矛盾"),
    ("allocation", "资源"),
    ("actions", "怎么做"),
    ("projection", "做完怎样"),
    ("transition", "转移"),
    ("review", "回头看"),
    ("risks", "注意事项"),
]


def load_request(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as fh:
        return json.load(fh)


def slugify(value: str) -> str:
    raw = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff-]+", "-", value.strip())
    raw = re.sub(r"-+", "-", raw).strip("-")
    return raw or "crux-report"


def clamp(value: float, low: float = 0.0, high: float = 1.0) -> float:
    return max(low, min(high, value))


def as_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None:
            return default
        return float(value)
    except (TypeError, ValueError):
        return default


def pct(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{clamp(float(value)) * 100:.1f}%"


def score(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{float(value):.2f}"


def display_text(value: Any) -> str:
    text = str(value if value is not None else "-")
    replacements = [
        ("监控阈值", "触发条件"),
        ("概率推演", "结果估计"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def md_escape(value: Any) -> str:
    return display_text(value).replace("|", "\\|").replace("\n", " ")


def html_text(value: Any) -> str:
    return escape(display_text(value))


def source_label(value: str | None) -> str:
    return SOURCE_LABELS.get(value or "assumed", value or "假设")


def confidence_label(value: str | None) -> str:
    return CONFIDENCE_LABELS.get(value or "medium", value or "中")


def plain_type_label(value: str | None) -> str:
    labels = {
        "value": "价值没闭上",
        "resource": "资源卡住",
        "process": "流程/排序问题",
        "market": "市场/成交问题",
        "goal": "目标拉扯",
        "relationship": "关系/信任问题",
        "information": "信息不够",
        "risk": "风险约束",
    }
    return labels.get(value or "", value or "-")


def score_contradiction(item: dict[str, Any]) -> float:
    scores = item.get("scores") or {}
    total = 0.0
    for key, weight in WEIGHTS.items():
        total += clamp(as_float(scores.get(key), 0.0), 0.0, 5.0) * weight
    return total


def contradiction_rows(request: dict[str, Any]) -> list[dict[str, Any]]:
    rows = []
    for item in request.get("contradictions") or []:
        row = dict(item)
        row["weighted_total"] = score_contradiction(item)
        rows.append(row)
    rows.sort(key=lambda row: row["weighted_total"], reverse=True)
    return rows


def string_list(value: Any, fallback: list[str]) -> list[str]:
    if isinstance(value, list):
        items = [str(item).strip() for item in value if str(item).strip()]
        return items or fallback
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return fallback


def fact_rows(request: dict[str, Any]) -> list[dict[str, str]]:
    rows = []
    for fact in request.get("facts") or []:
        rows.append(
            {
                "text": str(fact.get("text", "")),
                "source_type": str(fact.get("source_type", "assumed")),
                "note": str(fact.get("note", "")),
            }
        )
    return rows


def action_rows(request: dict[str, Any]) -> list[dict[str, Any]]:
    rows = []
    for item in request.get("actions") or []:
        completion = clamp(as_float(item.get("completion_probability"), 0.65))
        uplift = clamp(as_float(item.get("estimated_uplift"), 0.0))
        discount = clamp(as_float(item.get("dependency_discount"), 1.0))
        expected = completion * uplift * discount
        row = dict(item)
        row["completion_probability"] = completion
        row["estimated_uplift"] = uplift
        row["dependency_discount"] = discount
        row["expected_uplift"] = expected
        rows.append(row)
    return rows


def risk_rows(request: dict[str, Any]) -> list[dict[str, Any]]:
    rows = []
    for item in (request.get("projection") or {}).get("risk_drag") or []:
        probability = clamp(as_float(item.get("probability"), 0.0))
        impact = clamp(as_float(item.get("impact"), 0.0))
        row = dict(item)
        row["probability"] = probability
        row["impact"] = impact
        row["expected_drag"] = probability * impact
        rows.append(row)
    return rows


def build_projection(request: dict[str, Any], actions: list[dict[str, Any]], risks: list[dict[str, Any]]) -> dict[str, Any]:
    projection = request.get("projection") or {}
    baseline = clamp(as_float(projection.get("baseline_success_probability"), 0.25))
    expected_action_uplift = sum(action["expected_uplift"] for action in actions)
    full_uplift = sum(action["estimated_uplift"] * action["dependency_discount"] for action in actions)
    expected_risk_drag = sum(risk["expected_drag"] for risk in risks)
    recommended = clamp(baseline + expected_action_uplift - expected_risk_drag)
    high_quality = clamp(baseline + 1.15 * full_uplift - 0.6 * expected_risk_drag)
    low_execution = clamp(baseline + 0.35 * expected_action_uplift - 1.1 * expected_risk_drag)
    status_quo = clamp(baseline - 0.5 * expected_risk_drag)
    scenarios = [
        {
            "name": "维持现状",
            "probability": status_quo,
            "reading": "不集中处理关键卡点，只做零散优化，结果大概率继续被原来的卡点限制。",
        },
        {
            "name": "按建议执行",
            "probability": recommended,
            "reading": "完成关键动作并获得反馈后，目标达成概率的中性估计。",
        },
        {
            "name": "高质量执行",
            "probability": high_quality,
            "reading": "动作完成质量高、依赖少、风险被压住时的上行情景。",
        },
        {
            "name": "执行不足或受阻",
            "probability": low_execution,
            "reading": "动作完成不充分、并行干扰持续或风险变大时的下行情景。",
        },
    ]
    return {
        "baseline": baseline,
        "expected_action_uplift": expected_action_uplift,
        "expected_risk_drag": expected_risk_drag,
        "recommended_probability": recommended,
        "scenarios": scenarios,
        "sensitivity_notes": projection.get("sensitivity_notes") or [],
    }


def build_analysis_logic(request: dict[str, Any], rows: list[dict[str, Any]], principal: dict[str, Any]) -> dict[str, Any]:
    logic = request.get("analysis_logic") or {}
    situation = request.get("situation") or {}
    facts = fact_rows(request)
    visible_fallback = [fact["text"] for fact in facts[:4]]
    if not visible_fallback:
        visible_fallback = principal.get("evidence") or [situation.get("summary") or principal.get("name", "-")]
    hidden_fallback = [
        f"{principal.get('name', '-')}：它更像同时解释多个看得见问题的根部变量。"
    ]
    method_note = logic.get("method_note") or (
        "矛盾论提醒我们：复杂局面里矛盾很多，但当前阶段会有一种主要矛盾起领导和决定作用。"
        "所以这里不是按谁最吵、谁最急排序，而是看谁最能决定目标、牵动其他问题、占用稀缺资源。"
    )
    first_principles_question = logic.get("first_principles_question") or (
        "如果只能改变一个根部变量，哪一个会让多个看得见的问题一起变轻？"
    )
    higher_level_logic = logic.get("higher_level_logic") or (
        "上升一层看，不只比较用户描述里的表面问题，而是追问这些问题为什么会反复出现、"
        "共同消耗哪一种稀缺资源，以及哪个变量改变后能带动多个次要矛盾下降。"
    )
    why_not_surface = logic.get("why_not_surface") or (
        "看得见的问题需要处理，但如果它们只是同一个上游限制的表现，先逐个处理会平均用力。"
        f"按当前证据，{principal.get('name', '-')}更能解释其他问题为什么反复出现。"
    )
    ranking_rule = logic.get("ranking_rule") or (
        "先列出候选矛盾，再按目标影响、因果牵引、阶段紧迫、资源约束、可改变性、风险外溢和证据强度排序。"
        "如果两个候选接近，优先选择更上游、能释放最稀缺资源的那个。"
    )
    return {
        "method_note": method_note,
        "first_principles_question": first_principles_question,
        "higher_level_logic": higher_level_logic,
        "visible_problems": string_list(logic.get("visible_problems"), visible_fallback),
        "hidden_root_variables": string_list(logic.get("hidden_root_variables"), hidden_fallback),
        "why_not_surface": why_not_surface,
        "ranking_rule": ranking_rule,
        "candidate_count": len(rows),
    }


def shorten(value: Any, length: int = 28) -> str:
    text = display_text(value)
    return text if len(text) <= length else text[: length - 1] + "…"


def wrap_text(value: Any, length: int = 12, max_lines: int = 2) -> list[str]:
    text = display_text(value).replace("\n", " ").strip()
    if not text:
        return ["-"]
    lines: list[str] = []
    while text and len(lines) < max_lines:
        if len(lines) == max_lines - 1 and len(text) > length:
            lines.append(text[: max(1, length - 1)] + "…")
            break
        lines.append(text[:length])
        text = text[length:]
    return lines or ["-"]


def build_decision_model(request: dict[str, Any], rows: list[dict[str, Any]], principal: dict[str, Any]) -> dict[str, Any]:
    explicit = request.get("decision_model") or {}
    if explicit.get("candidates"):
        return explicit
    prior = 1 / max(len(rows), 1)
    candidates = []
    total_raw = 0.0
    for row in rows:
        scores = row.get("scores") or {}
        likelihood = clamp(row["weighted_total"] / 5)
        evidence_count = len(row.get("evidence") or [])
        assumption_count = as_float(row.get("assumption_count"), max(1, 4 - min(evidence_count, 4)))
        parsimony = clamp(1 - assumption_count * 0.08, 0.55, 1.0)
        leverage = clamp(as_float(scores.get("causal_leverage"), 0) / 5)
        raw = prior * (0.68 * likelihood + 0.32 * leverage) * parsimony
        total_raw += raw
        candidates.append(
            {
                "name": row.get("name", "-"),
                "prior": prior,
                "likelihood": likelihood,
                "causal_leverage": leverage,
                "assumption_count": assumption_count,
                "parsimony": parsimony,
                "posterior_raw": raw,
            }
        )
    for item in candidates:
        item["posterior"] = item["posterior_raw"] / total_raw if total_raw else 0
    candidates.sort(key=lambda item: item["posterior"], reverse=True)
    return {
        "bayesian_note": explicit.get("bayesian_note")
        or "贝叶斯更新：先把每个候选矛盾当作一个解释假设，再用目标影响、因果牵引、资源约束和证据强度更新可信度。",
        "ockham_note": explicit.get("ockham_note")
        or "奥卡姆剃刀：当两个候选解释力接近时，优先选择能解释更多表面问题、但额外假设更少的那个。",
        "selected": principal.get("name", "-"),
        "candidates": candidates,
    }


def build_resource_allocation(request: dict[str, Any]) -> dict[str, Any]:
    explicit = request.get("resource_allocation") or {}
    if explicit.get("items"):
        return explicit
    return {
        "title": "时间、精力、资源倾斜",
        "unit": "%",
        "note": "把负责人和团队的高杠杆资源从低价值救火中抽出来，转向能改变根部变量的动作。",
        "items": [
            {"label": "救火/临时响应", "current": 45, "recommended": 20},
            {"label": "沟通协调", "current": 25, "recommended": 15},
            {"label": "流程止血", "current": 20, "recommended": 25},
            {"label": "找人/系统建设", "current": 10, "recommended": 40},
        ],
    }


def build_stage_transition(request: dict[str, Any], principal: dict[str, Any], secondary: list[dict[str, Any]]) -> dict[str, Any]:
    explicit = request.get("stage_transition") or {}
    if explicit.get("stages"):
        return explicit
    next_items = secondary[:3] or [principal]
    return {
        "title": "主要矛盾动态转移",
        "note": "主要矛盾不是永久标签。当前主矛盾缓解后，下一阶段要重新看哪个矛盾开始决定局面。",
        "stages": [
            {
                "stage": "当前阶段",
                "principal": principal.get("name", "-"),
                "trigger": "当前目标窗口内最决定结果。",
            }
        ]
        + [
            {
                "stage": f"后续候选 {idx + 1}",
                "principal": row.get("name", "-"),
                "trigger": row.get("monitor_trigger") or row.get("reversal_condition") or "复盘指标触发后重新判断。",
            }
            for idx, row in enumerate(next_items)
        ],
    }


def build_visuals(request: dict[str, Any]) -> dict[str, Any]:
    explicit = request.get("visuals") or {}
    charts = explicit.get("charts") or []
    required = [
        ("analysis_flow", "分析流程图"),
        ("iceberg", "冰山模型"),
        ("decision_matrix", "矛盾候选决策矩阵"),
        ("resource_allocation", "时间/精力/资源倾斜图"),
        ("stage_transition", "动态阶段迁移图"),
    ]
    existing = {chart.get("type") for chart in charts if isinstance(chart, dict)}
    for kind, title in required:
        if kind not in existing:
            charts.append({"type": kind, "title": title})
    return {
        "summary": explicit.get("summary") or "用图表把表面问题、隐藏变量、候选矛盾、资源倾斜和动态转移串起来。",
        "charts": charts,
    }


def build_report(request: dict[str, Any]) -> dict[str, Any]:
    rows = contradiction_rows(request)
    if not rows:
        raise ValueError("request.contradictions must contain at least one item")
    principal = rows[0]
    secondary = rows[1:]
    actions = action_rows(request)
    risks = risk_rows(request)
    projection = build_projection(request, actions, risks)
    current_state_clarity = assess_current_state_clarity(request)
    analysis_logic = build_analysis_logic(request, rows, principal)
    decision_model = build_decision_model(request, rows, principal)
    resource_allocation = build_resource_allocation(request)
    stage_transition = build_stage_transition(request, principal, secondary)
    visuals = build_visuals(request)
    confidence = request.get("confidence") or "medium"
    principal_score = principal["weighted_total"]
    first_action = actions[0]["action"] if actions else "先补齐目标、阶段、资源和事实证据。"
    if principal_score < 3.5:
        one_sentence = f"现在还不能稳定判断最关键卡点，先验证：{principal.get('name', '-')}"
    else:
        one_sentence = f"现在最该先解决的是：{principal.get('name', '-')}"
    display_title = request.get("plain_title") or one_sentence
    report = {
        "title": display_title,
        "source_title": request.get("title") or display_title,
        "slug": request.get("slug") or slugify(request.get("title") or "crux-report"),
        "case_type": request.get("case_type", "general"),
        "analysis_date": request.get("analysis_date") or date.today().isoformat(),
        "confidence": confidence,
        "situation": request.get("situation") or {},
        "facts": fact_rows(request),
        "contradictions": rows,
        "principal": principal,
        "secondary": secondary,
        "actions": actions,
        "risks": risks,
        "projection": projection,
        "analysis_logic": analysis_logic,
        "decision_model": decision_model,
        "resource_allocation": resource_allocation,
        "stage_transition": stage_transition,
        "visuals": visuals,
        "current_state_clarity": current_state_clarity,
        "review": request.get("review") or {},
        "warnings": request.get("warnings") or [],
        "summary": {
            "one_sentence": one_sentence,
            "principal_aspect": principal.get("principal_aspect") or "需要继续追问以判断主要方面。",
            "confidence_label": confidence_label(confidence),
            "principal_score": principal_score,
            "first_action": first_action,
            "recommended_probability": projection["recommended_probability"],
            "current_state_score": current_state_clarity["score"],
            "current_state_level": current_state_clarity["clarity_level"],
        },
    }
    return report


def list_md(items: list[Any]) -> str:
    if not items:
        return "- -"
    return "\n".join(f"- {md_escape(item)}" for item in items)


def md_table(headers: list[str], rows: list[list[Any]]) -> str:
    output = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        output.append("| " + " | ".join(md_escape(value) for value in row) + " |")
    return "\n".join(output)


def render_markdown(report: dict[str, Any]) -> str:
    situation = report["situation"]
    principal = report["principal"]
    projection = report["projection"]
    analysis_logic = report["analysis_logic"]
    decision_model = report["decision_model"]
    allocation = report["resource_allocation"]
    stage_transition = report["stage_transition"]
    lines = [
        f"# {display_text(report['title'])}",
        "",
        f"> {display_text(report['summary']['one_sentence'])}。现在主导局面的因素：{display_text(report['summary']['principal_aspect'])}。",
        "",
        "## 先看结论",
        "",
        md_table(
            ["最该先解决", "为什么是它", "把握", "做完后大概到哪", "第一步"],
            [[
                principal.get("name", "-"),
                principal.get("why_principal") or report["summary"]["principal_aspect"],
                report["summary"]["confidence_label"],
                pct(report["summary"]["recommended_probability"]),
                report["summary"]["first_action"],
            ]],
        ),
        "",
        "## 先把现状说清楚",
        "",
        display_text(situation.get("summary", "-")),
        "",
        md_table(
            ["目标", "期限", "成功标准", "阶段"],
            [[situation.get("goal", "-"), situation.get("deadline", "-"), situation.get("success_metric", "-"), situation.get("stage", "-")]],
        ),
        "",
        "### 资源、约束与相关方",
        "",
        md_table(
            ["资源", "约束", "相关方"],
            [[", ".join(situation.get("resources") or ["-"]), ", ".join(situation.get("constraints") or ["-"]), ", ".join(situation.get("stakeholders") or ["-"])]],
        ),
        "",
        "## 现状够不够清楚",
        "",
        f"```text\n{display_text(report['current_state_clarity']['snapshot'])}\n```",
        "",
        md_table(
            ["维度", "得分"],
            [[CLARITY_DIMENSION_LABELS.get(key, key), f"{value:.0f}"] for key, value in report["current_state_clarity"]["dimension_scores"].items()],
        ),
        "",
        "### 如果还不够清楚，先问这些",
        "",
        list_md(report["current_state_clarity"].get("questions") or ["现状已经够用，下一步只需要让用户确认上面的复述是否准确。"]),
        "",
        "### 哪些是事实，哪些还只是判断",
        "",
        md_table(
            ["事实", "类型", "备注"],
            [[fact["text"], source_label(fact["source_type"]), fact.get("note") or "-"] for fact in report["facts"]],
        ),
        "",
        "## 一张图看懂：从表象到主要矛盾",
        "",
        display_text(report["visuals"]["summary"]),
        "",
        "### 分析流程图",
        "",
        "用户描述 -> 追问补全现状 -> 看得见的问题 -> 上升一层 -> 看不见的根部变量 -> 贝叶斯更新 -> 奥卡姆剃刀 -> 主要矛盾/次要矛盾 -> 行动与复盘",
        "",
        "### 冰山模型",
        "",
        md_table(
            ["层级", "内容"],
            [
                ["水面上：看得见的问题", "；".join(analysis_logic["visible_problems"][:4])],
                ["水面下：看不见的根部变量", "；".join(analysis_logic["hidden_root_variables"][:4])],
                ["最底层：当前主要矛盾", principal.get("name", "-")],
            ],
        ),
        "",
        "### 贝叶斯更新与奥卡姆剃刀",
        "",
        display_text(decision_model["bayesian_note"]),
        "",
        display_text(decision_model["ockham_note"]),
        "",
        md_table(
            ["候选矛盾", "更新后可信度", "解释力", "简洁度"],
            [
                [row["name"], pct(row.get("posterior")), pct(row.get("likelihood")), pct(row.get("parsimony"))]
                for row in decision_model["candidates"][:5]
            ],
        ),
        "",
        "### 矛盾候选决策矩阵",
        "",
        md_table(
            ["候选矛盾", "目标影响", "因果牵引", "资源约束", "证据强度"],
            [
                [
                    row.get("name", "-"),
                    score((row.get("scores") or {}).get("goal_impact")),
                    score((row.get("scores") or {}).get("causal_leverage")),
                    score((row.get("scores") or {}).get("resource_constraint")),
                    score((row.get("scores") or {}).get("evidence_strength")),
                ]
                for row in report["contradictions"][:5]
            ],
        ),
        "",
        "## 主要矛盾判断过程",
        "",
        display_text(analysis_logic["method_note"]),
        "",
        f"第一性原理追问：{display_text(analysis_logic['first_principles_question'])}",
        "",
        "### 看得见的问题",
        "",
        list_md(analysis_logic["visible_problems"]),
        "",
        "### 上升一层：找看不见的根部变量",
        "",
        display_text(analysis_logic["higher_level_logic"]),
        "",
        list_md(analysis_logic["hidden_root_variables"]),
        "",
        "### 为什么不是先处理表面问题",
        "",
        display_text(analysis_logic["why_not_surface"]),
        "",
        "### 候选矛盾怎么排序",
        "",
        display_text(analysis_logic["ranking_rule"]),
        "",
        md_table(
            ["排序", "可能的卡点", "大致类型", "优先级", "依据"],
            [
                [idx + 1, row.get("name", "-"), plain_type_label(row.get("type")), f"{row['weighted_total']:.2f}/5", "；".join(row.get("evidence") or ["-"])]
                for idx, row in enumerate(report["contradictions"])
            ],
        ),
        "",
        "## 主要矛盾（最关键的卡点）",
        "",
        f"- 主要矛盾：{display_text(principal.get('name', '-'))}",
        f"- 主要方面（现在最影响局面的一侧）：{display_text(principal.get('principal_aspect') or '需要继续追问')}",
        f"- 为什么先处理它：{display_text(principal.get('why_principal') or '按现有事实看，它最影响目标推进。')}",
        f"- 什么情况说明判断要改：{display_text(principal.get('reversal_condition') or '当关键指标改善但结果仍未变化时，需要重新判断。')}",
        "",
        "## 次要矛盾（先不主攻，但要盯住）",
        "",
        md_table(
            ["次要矛盾", "为什么暂时不主攻", "什么时候要处理"],
            [[row.get("name", "-"), row.get("defer_reason") or "当前不是资源投放主线。", row.get("monitor_trigger") or "阶段变化时再看。"] for row in report["secondary"]],
        ),
        "",
        "## 时间、精力、资源应该怎么重新分配",
        "",
        display_text(allocation.get("note", "-")),
        "",
        md_table(
            ["资源项", "当前分配", "建议分配", "变化方向"],
            [
                [
                    item.get("label", "-"),
                    f"{as_float(item.get('current')):.0f}{allocation.get('unit', '%')}",
                    f"{as_float(item.get('recommended')):.0f}{allocation.get('unit', '%')}",
                    f"{as_float(item.get('recommended')) - as_float(item.get('current')):+.0f}{allocation.get('unit', '%')}",
                ]
                for item in allocation.get("items", [])
            ],
        ),
        "",
        "## 接下来怎么做",
        "",
        md_table(
            ["动作", "谁来做", "什么时候", "要用什么", "做到什么算完成", "做成把握", "预计帮助"],
            [
                [
                    row.get("action", "-"),
                    row.get("owner", "-"),
                    row.get("deadline", "-"),
                    row.get("resource", "-"),
                    row.get("metric", "-"),
                    pct(row["completion_probability"]),
                    pct(row["expected_uplift"]),
                ]
                for row in report["actions"]
            ],
        ),
        "",
        "## 做完以后可能怎样",
        "",
        md_table(
            ["不专门处理时", "这些动作可能带来的帮助", "风险可能拖后腿", "按建议做后的估计"],
            [[pct(projection["baseline"]), pct(projection["expected_action_uplift"]), pct(projection["expected_risk_drag"]), pct(projection["recommended_probability"])]],
        ),
        "",
        md_table(
            ["情况", "大概概率", "怎么理解"],
            [[item["name"], pct(item["probability"]), item["reading"]] for item in projection["scenarios"]],
        ),
        "",
        "### 哪些变化会影响这个估计",
        "",
        list_md(projection.get("sensitivity_notes") or ["行动后概率依赖用户补充证据和执行反馈。"]),
        "",
        "## 主要矛盾什么时候会转移",
        "",
        display_text(stage_transition.get("note", "-")),
        "",
        md_table(
            ["阶段", "可能主导的矛盾", "触发条件"],
            [[item.get("stage", "-"), item.get("principal", "-"), item.get("trigger", "-")] for item in stage_transition.get("stages", [])],
        ),
        "",
        "## 什么时候回头看",
        "",
        f"- 复盘时间：{report['review'].get('review_time', '7到30天后复盘')}",
        "",
        "### 看到这些信号就重新判断",
        "",
        list_md(report["review"].get("review_conditions") or []),
        "",
        "### 下次要补的证据",
        "",
        list_md(report["review"].get("evidence_to_collect") or []),
        "",
        "## 注意事项",
        "",
        list_md(report["warnings"] or ["本报告只用于结构化诊断和行动辅助，不替代专业意见。"]),
        "",
    ]
    return "\n".join(lines)


def html_list(items: list[Any]) -> str:
    if not items:
        return "<ul><li>-</li></ul>"
    return "<ul>" + "".join(f"<li>{html_text(item)}</li>" for item in items) + "</ul>"


def html_table(headers: list[str], rows: list[list[Any]], top_first: bool = False) -> str:
    body = []
    for idx, row in enumerate(rows):
        klass = ' class="top-row"' if top_first and idx == 0 else ""
        body.append("<tr" + klass + ">" + "".join(f"<td>{html_text(value)}</td>" for value in row) + "</tr>")
    return (
        "<table><thead><tr>"
        + "".join(f"<th>{html_text(header)}</th>" for header in headers)
        + "</tr></thead><tbody>"
        + "".join(body)
        + "</tbody></table>"
    )


def probability_bar(value: float) -> str:
    width = clamp(value) * 100
    return f'<div class="probability-bar"><div class="probability-fill" style="width:{width:.1f}%"></div></div>'


SVG_TEXT_STYLE = {
    "chart-text": ('#141413', 17),
    "chart-text-strong": ('#ffffff', 17),
    "chart-note": ('#6b6a64', 14),
    "chart-kicker": ('#1B365D', 15),
}


def svg_text(value: Any, x: float, y: float, klass: str = "chart-text", anchor: str = "middle") -> str:
    fill, size = SVG_TEXT_STYLE.get(klass, SVG_TEXT_STYLE["chart-text"])
    return (
        f'<text class="{klass}" x="{x:.1f}" y="{y:.1f}" text-anchor="{anchor}" '
        f'fill="{fill}" font-size="{size}" font-family="Songti SC, STSong, Georgia, serif">'
        f'{html_text(shorten(value, 22))}</text>'
    )


def svg_multiline(
    value: Any,
    x: float,
    y: float,
    klass: str = "chart-text",
    anchor: str = "middle",
    length: int = 12,
    max_lines: int = 2,
    line_height: int = 22,
) -> str:
    fill, size = SVG_TEXT_STYLE.get(klass, SVG_TEXT_STYLE["chart-text"])
    lines = wrap_text(value, length=length, max_lines=max_lines)
    spans = []
    for idx, line in enumerate(lines):
        dy = 0 if idx == 0 else line_height
        spans.append(f'<tspan x="{x:.1f}" dy="{dy}">{html_text(line)}</tspan>')
    return (
        f'<text class="{klass}" x="{x:.1f}" y="{y:.1f}" text-anchor="{anchor}" '
        f'fill="{fill}" font-size="{size}" font-family="Songti SC, STSong, Georgia, serif">'
        f'{"".join(spans)}</text>'
    )


def image_data_uri(path: Path, mime_type: str | None = None) -> str:
    if mime_type is None:
        mime_type = "image/png" if path.suffix.lower() == ".png" else "image/jpeg"
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime_type};base64,{data}"


def html_compact_items(items: list[Any], limit: int = 3, length: int = 34) -> str:
    selected = items[:limit] or ["-"]
    return "<ul>" + "".join(f"<li>{html_text(shorten(item, length))}</li>" for item in selected) + "</ul>"


def svg_analysis_flow(report: dict[str, Any]) -> str:
    steps = ["用户描述", "追问补全", "看得见的问题", "上升一层", "贝叶斯更新", "奥卡姆剃刀", "主要矛盾"]
    parts = ['<svg id="chart-analysis-flow" class="chart-svg" viewBox="0 0 960 230" role="img" aria-label="分析流程图">']
    for idx, label in enumerate(steps):
        x = 24 + idx * 132
        parts.append(f'<rect class="chart-node" x="{x}" y="80" width="108" height="54" rx="8" fill="#EEF2F7" stroke="#cad4e1"></rect>')
        parts.append(svg_text(label, x + 54, 112))
        if idx < len(steps) - 1:
            line_x = x + 108
            parts.append(f'<path class="chart-arrow" d="M {line_x + 6} 107 L {line_x + 18} 107" fill="none" stroke="#1B365D" stroke-width="2"></path>')
            parts.append(f'<path class="chart-arrow-head" d="M {line_x + 18} 107 l -6 -5 v 10 z" fill="#1B365D"></path>')
    parts.append(svg_text("从用户能看见的问题，走向更上游的主要矛盾", 480, 178, "chart-note"))
    parts.append("</svg>")
    return "".join(parts)


def svg_iceberg(report: dict[str, Any]) -> str:
    logic = report["analysis_logic"]
    visible = [item.replace("看得见的问题：", "") for item in logic["visible_problems"][:3]]
    hidden = [item.replace("看不见的根部变量：", "") for item in logic["hidden_root_variables"][:3]]
    principal = report["principal"].get("name", "-")
    photo_uri = image_data_uri(ICEBERG_PHOTO_PATH)
    return f"""
<figure id="chart-iceberg" class="iceberg-model iceberg-photo-model" role="img" aria-label="冰山模型：从看得见的问题到看不见的根部变量">
  <div class="iceberg-visual">
    <img class="iceberg-photo" src="{photo_uri}" alt="真实冰山水上和水下照片">
    <div class="iceberg-waterline"><span>水面线</span></div>
  </div>
  <div class="iceberg-layers">
    <section class="iceberg-card iceberg-card-surface">
      <div class="iceberg-label-title">水面上：看得见的问题</div>
      {html_compact_items(visible, 3, 34)}
    </section>
    <section class="iceberg-card iceberg-card-hidden">
      <div class="iceberg-label-title">水面下：看不见的根部变量</div>
      {html_compact_items(hidden, 3, 36)}
    </section>
    <section class="iceberg-card iceberg-card-principal">
      <div class="iceberg-label-title">当前主要矛盾</div>
      <p>{html_text(shorten(principal, 46))}</p>
    </section>
  </div>
  <figcaption class="iceberg-source">冰山图片：仓库内置素材</figcaption>
</figure>
""".strip()


def svg_decision_matrix(report: dict[str, Any]) -> str:
    rows = report["contradictions"][:5]
    parts = ['<svg id="chart-decision-matrix" class="chart-svg chart-svg-tall" viewBox="0 0 960 520" role="img" aria-label="矛盾候选决策矩阵">']
    plot_x, plot_y, plot_w, plot_h = 88, 86, 570, 310
    parts.append(f'<rect class="chart-plot-bg" x="{plot_x}" y="{plot_y}" width="{plot_w}" height="{plot_h}" rx="8" fill="#fffdf8" stroke="#e8e6dc"></rect>')
    for idx in range(1, 5):
        x = plot_x + idx * plot_w / 5
        y = plot_y + idx * plot_h / 5
        parts.append(f'<path class="chart-grid-line" d="M{x:.1f} {plot_y} L{x:.1f} {plot_y + plot_h} M{plot_x} {y:.1f} L{plot_x + plot_w} {y:.1f}" fill="none" stroke="#e8e6dc" stroke-width="1"></path>')
    parts.append(f'<path class="chart-axis" d="M{plot_x} {plot_y + plot_h} L{plot_x + plot_w} {plot_y + plot_h} M{plot_x} {plot_y + plot_h} L{plot_x} {plot_y}" fill="none" stroke="#1B365D" stroke-width="2"></path>')
    parts.append(svg_text("因果牵引力", plot_x + plot_w - 4, plot_y + plot_h + 42, "chart-note", "end"))
    parts.append(svg_text("目标影响力", plot_x - 12, plot_y - 20, "chart-note", "middle"))
    jitter = [(-20, -18), (22, 16), (-18, 22), (18, -24), (0, 0)]
    for idx, row in enumerate(rows):
        scores = row.get("scores") or {}
        raw_x = plot_x + clamp(as_float(scores.get("causal_leverage")) / 5) * (plot_w - 42)
        raw_y = plot_y + plot_h - clamp(as_float(scores.get("goal_impact")) / 5) * (plot_h - 42)
        dx, dy = jitter[idx % len(jitter)]
        x = clamp(raw_x + dx, plot_x + 28, plot_x + plot_w - 28)
        y = clamp(raw_y + dy, plot_y + 28, plot_y + plot_h - 28)
        radius = 14 + clamp(as_float(scores.get("resource_constraint")) / 5) * 14
        klass = "matrix-dot-primary" if idx == 0 else "matrix-dot"
        fill = "#1B365D" if idx == 0 else "#d6d1c3"
        stroke = "#1B365D" if idx == 0 else "#6b6a64"
        width = 2 if idx == 0 else 1
        parts.append(f'<circle class="{klass}" cx="{x:.1f}" cy="{y:.1f}" r="{radius:.1f}" fill="{fill}" stroke="{stroke}" stroke-width="{width}"></circle>')
        parts.append(svg_text(str(idx + 1), x, y + 5, "chart-text-strong" if idx == 0 else "chart-text"))
    parts.append('<rect class="chart-legend-box" x="696" y="86" width="214" height="310" rx="8" fill="#faf9f5" stroke="#e8e6dc"></rect>')
    parts.append(svg_text("候选矛盾", 718, 122, "chart-kicker", "start"))
    for idx, row in enumerate(rows):
        y = 154 + idx * 48
        index_class = "chart-legend-index-primary" if idx == 0 else "chart-legend-index"
        fill = "#1B365D" if idx == 0 else "#d6d1c3"
        stroke = "#1B365D" if idx == 0 else "#6b6a64"
        parts.append(f'<circle class="{index_class}" cx="718" cy="{y - 6}" r="11" fill="{fill}" stroke="{stroke}"></circle>')
        parts.append(svg_text(str(idx + 1), 718, y - 2, "chart-text-strong"))
        parts.append(svg_multiline(row.get("name", "-"), 742, y - 10, "chart-note", "start", length=14, max_lines=2, line_height=18))
    parts.append(svg_text("越靠右上，越应优先处理；圆越大，越消耗稀缺资源。", 480, 462, "chart-note"))
    parts.append("</svg>")
    return "".join(parts)


def svg_resource_allocation(report: dict[str, Any]) -> str:
    allocation = report["resource_allocation"]
    items = allocation.get("items", [])[:5]
    parts = ['<svg id="chart-resource-allocation" class="chart-svg" viewBox="0 0 960 420" role="img" aria-label="资源倾斜图">']
    max_value = max([as_float(item.get("current")) for item in items] + [as_float(item.get("recommended")) for item in items] + [1])
    for idx, item in enumerate(items):
        y = 70 + idx * 68
        current_w = as_float(item.get("current")) / max_value * 300
        recommended_w = as_float(item.get("recommended")) / max_value * 300
        parts.append(svg_text(item.get("label", "-"), 74, y + 20, "chart-note", "start"))
        parts.append(f'<rect class="bar-current" x="250" y="{y}" width="{current_w:.1f}" height="18" rx="4" fill="#d6d1c3"></rect>')
        parts.append(f'<rect class="bar-recommended" x="250" y="{y + 26}" width="{recommended_w:.1f}" height="18" rx="4" fill="#1B365D"></rect>')
        parts.append(svg_text(f"当前 {as_float(item.get('current')):.0f}{allocation.get('unit', '%')}", 570, y + 15, "chart-note", "start"))
        parts.append(svg_text(f"建议 {as_float(item.get('recommended')):.0f}{allocation.get('unit', '%')}", 570, y + 41, "chart-note", "start"))
    parts.append('<rect class="bar-current" x="700" y="338" width="18" height="12" fill="#d6d1c3"></rect>')
    parts.append(svg_text("当前", 746, 349, "chart-note", "start"))
    parts.append('<rect class="bar-recommended" x="700" y="362" width="18" height="12" fill="#1B365D"></rect>')
    parts.append(svg_text("建议", 746, 373, "chart-note", "start"))
    parts.append("</svg>")
    return "".join(parts)


def svg_stage_transition(report: dict[str, Any]) -> str:
    stages = report["stage_transition"].get("stages", [])[:4]
    parts = ['<svg id="chart-stage-transition" class="chart-svg chart-svg-tall" viewBox="0 0 960 620" role="img" aria-label="动态阶段迁移图">']
    parts.append('<path class="stage-spine" d="M118 88 L118 524" fill="none" stroke="#1B365D" stroke-width="2"></path>')
    for idx, item in enumerate(stages):
        y = 72 + idx * 124
        parts.append(f'<circle class="stage-dot" cx="118" cy="{y + 45}" r="18" fill="#1B365D" stroke="#1B365D"></circle>')
        parts.append(svg_text(str(idx + 1), 118, y + 51, "chart-text-strong"))
        parts.append(f'<rect class="stage-card" x="164" y="{y}" width="676" height="98" rx="8" fill="#faf9f5" stroke="#e8e6dc"></rect>')
        parts.append(svg_text(item.get("stage", "-"), 190, y + 31, "chart-kicker", "start"))
        parts.append(svg_multiline(item.get("principal", "-"), 190, y + 58, "chart-text", "start", length=24, max_lines=1))
        parts.append(svg_multiline(item.get("trigger", "-"), 190, y + 82, "chart-note", "start", length=42, max_lines=1))
    parts.append(svg_text("阶段变化后，主要矛盾也要重新判断", 502, 578, "chart-note"))
    parts.append("</svg>")
    return "".join(parts)


def render_html(report: dict[str, Any]) -> str:
    css = CSS_PATH.read_text(encoding="utf-8")
    situation = report["situation"]
    principal = report["principal"]
    projection = report["projection"]
    analysis_logic = report["analysis_logic"]
    decision_model = report["decision_model"]
    allocation = report["resource_allocation"]
    stage_transition = report["stage_transition"]
    nav = "".join(f'<a href="#{anchor}">{label}</a>' for anchor, label in NAV_ITEMS)
    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html_text(report['title'])}</title>
  <meta name="author" content="Yao Crux Skill">
  <meta name="description" content="{html_text(report['summary']['one_sentence'])}">
  <meta name="generator" content="Yao Crux Skill">
  <style>{css}</style>
</head>
<body>
  <header class="topbar">
    <div class="topbar-inner">
      <div class="brand">
        <div class="brand-title">现状诊断报告</div>
        <div class="brand-subtitle">{html_text(report['analysis_date'])}</div>
      </div>
      <nav class="nav">{nav}</nav>
      <button class="print-button" onclick="window.print()">打印 / 存为 PDF</button>
    </div>
  </header>
  <main class="page">
    <section class="hero" id="summary">
      <div class="hero-grid">
        <div class="hero-main">
          <div class="eyebrow">先看结论</div>
          <h1>{html_text(report['title'])}</h1>
          <p class="lead">{html_text(report['summary']['one_sentence'])}。现在主导局面的因素：{html_text(report['summary']['principal_aspect'])}。</p>
        </div>
        <aside class="hero-action">
          <div class="metric-label">先做这一步</div>
          <p>{html_text(report['summary']['first_action'])}</p>
        </aside>
      </div>
      <div class="summary-strip">
        <div class="summary-item"><div class="metric-label">判断把握</div><div class="summary-value">{html_text(report['summary']['confidence_label'])}</div></div>
        <div class="summary-item"><div class="metric-label">现状清晰度</div><div class="summary-value">{html_text(report['summary']['current_state_score'])}/100</div></div>
        <div class="summary-item"><div class="metric-label">做完后估计</div><div class="summary-value">{pct(report['summary']['recommended_probability'])}</div>{probability_bar(report['summary']['recommended_probability'])}</div>
      </div>
    </section>

    <section class="section" id="situation">
      <div class="section-title"><div class="section-kicker">Context</div><h2>先把现状说清楚</h2></div>
      <div class="panel">
        <p>{html_text(situation.get('summary', '-'))}</p>
        {html_table(['目标', '期限', '成功标准', '阶段'], [[situation.get('goal', '-'), situation.get('deadline', '-'), situation.get('success_metric', '-'), situation.get('stage', '-')]])}
        {html_table(['资源', '约束', '相关方'], [[', '.join(situation.get('resources') or ['-']), ', '.join(situation.get('constraints') or ['-']), ', '.join(situation.get('stakeholders') or ['-'])]])}
      </div>
    </section>

    <section class="section" id="current-state">
      <div class="section-title"><div class="section-kicker">Clarify first</div><h2>现状够不够清楚</h2></div>
      <div class="panel">
        <div class="callout"><strong>能不能开始判断：</strong>{html_text('可以继续判断关键卡点' if report['current_state_clarity']['diagnosis_allowed'] else '还不能直接下结论，先追问')}</div>
        <p><strong>清晰度：</strong>{html_text(report['current_state_clarity']['score'])}/100（{html_text(CLARITY_LEVEL_LABELS.get(report['current_state_clarity']['clarity_level'], report['current_state_clarity']['clarity_level']))}）</p>
        <pre>{html_text(report['current_state_clarity']['snapshot'])}</pre>
        {html_table(['维度', '得分'], [[CLARITY_DIMENSION_LABELS.get(key, key), f"{value:.0f}"] for key, value in report['current_state_clarity']['dimension_scores'].items()])}
        <h3>如果还不够清楚，先问这些</h3>
        {html_list(report['current_state_clarity'].get('questions') or ['现状已经够用，下一步只需要让用户确认上面的复述是否准确。'])}
      </div>
      <div class="panel">
        <h3>哪些是事实，哪些还只是判断</h3>
        {html_table(['事实', '类型', '备注'], [[fact['text'], source_label(fact['source_type']), fact.get('note') or '-'] for fact in report['facts']])}
      </div>
    </section>

    <section class="section" id="visuals">
      <div class="section-title"><div class="section-kicker">Visual reasoning</div><h2>一张图看懂：从表象到主要矛盾</h2></div>
      <div class="panel">
        <p>{html_text(report['visuals']['summary'])}</p>
        <h3>分析流程图</h3>
        {svg_analysis_flow(report)}
        <h3>冰山模型</h3>
        {svg_iceberg(report)}
        <h3>贝叶斯更新与奥卡姆剃刀</h3>
        <p>{html_text(decision_model['bayesian_note'])}</p>
        <p>{html_text(decision_model['ockham_note'])}</p>
        {html_table(['候选矛盾', '更新后可信度', '解释力', '简洁度'], [[row['name'], pct(row.get('posterior')), pct(row.get('likelihood')), pct(row.get('parsimony'))] for row in decision_model['candidates'][:5]], top_first=True)}
        <h3>矛盾候选决策矩阵</h3>
        {svg_decision_matrix(report)}
      </div>
    </section>

    <section class="section" id="contradictions">
      <div class="section-title"><div class="section-kicker">Reasoning</div><h2>主要矛盾判断过程</h2></div>
      <div class="panel">
        <p>{html_text(analysis_logic['method_note'])}</p>
        <div class="callout"><strong>第一性原理追问：</strong>{html_text(analysis_logic['first_principles_question'])}</div>
        <h3>看得见的问题</h3>
        {html_list(analysis_logic['visible_problems'])}
        <h3>上升一层：找看不见的根部变量</h3>
        <p>{html_text(analysis_logic['higher_level_logic'])}</p>
        {html_list(analysis_logic['hidden_root_variables'])}
        <h3>为什么不是先处理表面问题</h3>
        <p>{html_text(analysis_logic['why_not_surface'])}</p>
        <h3>候选矛盾怎么排序</h3>
        <p>{html_text(analysis_logic['ranking_rule'])}</p>
        {html_table(['排序', '可能的卡点', '大致类型', '优先级', '依据'], [[idx + 1, row.get('name', '-'), plain_type_label(row.get('type')), f"{row['weighted_total']:.2f}/5", '；'.join(row.get('evidence') or ['-'])] for idx, row in enumerate(report['contradictions'])], top_first=True)}
      </div>
    </section>

    <section class="section" id="principal">
      <div class="section-title"><div class="section-kicker">Focus</div><h2>主要矛盾（最关键的卡点）</h2></div>
      <div class="panel">
        <div class="callout"><strong>主要矛盾：</strong>{html_text(principal.get('name', '-'))}</div>
        <p><strong>主要方面（现在最影响局面的一侧）：</strong>{html_text(principal.get('principal_aspect') or '需要继续追问')}</p>
        <p><strong>为什么先处理它：</strong>{html_text(principal.get('why_principal') or '按现有事实看，它最影响目标推进。')}</p>
        <p><strong>什么情况说明判断要改：</strong>{html_text(principal.get('reversal_condition') or '当关键指标改善但结果仍未变化时，需要重新判断。')}</p>
      </div>
    </section>

    <section class="section" id="secondary">
      <div class="section-title"><div class="section-kicker">Watchlist</div><h2>次要矛盾（先不主攻，但要盯住）</h2></div>
      <div class="panel">
        {html_table(['次要矛盾', '为什么暂时不主攻', '什么时候要处理'], [[row.get('name', '-'), row.get('defer_reason') or '当前不是资源投放主线。', row.get('monitor_trigger') or '阶段变化时再看。'] for row in report['secondary']])}
      </div>
    </section>

    <section class="section" id="allocation">
      <div class="section-title"><div class="section-kicker">Resource tilt</div><h2>时间、精力、资源应该怎么重新分配</h2></div>
      <div class="panel">
        <p>{html_text(allocation.get('note', '-'))}</p>
        {svg_resource_allocation(report)}
        {html_table(['资源项', '当前分配', '建议分配', '变化方向'], [[item.get('label', '-'), f"{as_float(item.get('current')):.0f}{allocation.get('unit', '%')}", f"{as_float(item.get('recommended')):.0f}{allocation.get('unit', '%')}", f"{as_float(item.get('recommended')) - as_float(item.get('current')):+.0f}{allocation.get('unit', '%')}"] for item in allocation.get('items', [])])}
      </div>
    </section>

    <section class="section" id="actions">
      <div class="section-title"><div class="section-kicker">Actions</div><h2>接下来怎么做</h2></div>
      <div class="panel">
        {html_table(['动作', '谁来做', '什么时候', '要用什么', '做到什么算完成', '做成把握', '预计帮助'], [[row.get('action', '-'), row.get('owner', '-'), row.get('deadline', '-'), row.get('resource', '-'), row.get('metric', '-'), pct(row['completion_probability']), pct(row['expected_uplift'])] for row in report['actions']])}
      </div>
    </section>

    <section class="section" id="projection">
      <div class="section-title"><div class="section-kicker">Projection</div><h2>做完以后可能怎样</h2></div>
      <div class="panel">
        {html_table(['不专门处理时', '这些动作可能带来的帮助', '风险可能拖后腿', '按建议做后的估计'], [[pct(projection['baseline']), pct(projection['expected_action_uplift']), pct(projection['expected_risk_drag']), pct(projection['recommended_probability'])]])}
        {html_table(['情况', '大概概率', '怎么理解'], [[item['name'], pct(item['probability']), item['reading']] for item in projection['scenarios']])}
        <h3>哪些变化会影响这个估计</h3>
        {html_list(projection.get('sensitivity_notes') or ['行动后概率依赖用户补充证据和执行反馈。'])}
      </div>
    </section>

    <section class="section" id="transition">
      <div class="section-title"><div class="section-kicker">Dynamic shift</div><h2>主要矛盾什么时候会转移</h2></div>
      <div class="panel">
        <p>{html_text(stage_transition.get('note', '-'))}</p>
        {svg_stage_transition(report)}
        {html_table(['阶段', '可能主导的矛盾', '触发条件'], [[item.get('stage', '-'), item.get('principal', '-'), item.get('trigger', '-')] for item in stage_transition.get('stages', [])])}
      </div>
    </section>

    <section class="section" id="review">
      <div class="section-title"><div class="section-kicker">Review</div><h2>什么时候回头看</h2></div>
      <div class="panel">
        <p><strong>复盘时间：</strong>{html_text(report['review'].get('review_time', '7到30天后复盘'))}</p>
        <h3>看到这些信号就重新判断</h3>
        {html_list(report['review'].get('review_conditions') or [])}
        <h3>下次要补的证据</h3>
        {html_list(report['review'].get('evidence_to_collect') or [])}
      </div>
    </section>

    <section class="section" id="risks">
      <div class="section-title"><div class="section-kicker">Boundaries</div><h2>注意事项</h2></div>
      <div class="panel">
        {html_list(report['warnings'] or ['本报告只用于结构化诊断和行动辅助，不替代专业意见。'])}
      </div>
    </section>
  </main>
</body>
</html>
"""
    return html


def add_docx_table(document: Any, headers: list[str], rows: list[list[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = display_text(value)


def add_docx_bullets(document: Any, items: list[Any]) -> None:
    for item in items or ["-"]:
        document.add_paragraph(display_text(item), style="List Bullet")


def write_docx(report: dict[str, Any], path: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Songti SC"
    styles["Normal"].font.size = Pt(10.5) if Pt else None
    document.add_heading(display_text(report["title"]), level=0)
    document.add_paragraph(f"{display_text(report['summary']['one_sentence'])}。现在主导局面的因素：{display_text(report['summary']['principal_aspect'])}。")
    document.add_heading("先看结论", level=1)
    add_docx_table(
        document,
        ["最该先解决", "为什么是它", "把握", "做完后大概到哪", "第一步"],
        [[
            report["principal"].get("name", "-"),
            report["principal"].get("why_principal") or report["summary"]["principal_aspect"],
            report["summary"]["confidence_label"],
            pct(report["summary"]["recommended_probability"]),
            report["summary"]["first_action"],
        ]],
    )
    situation = report["situation"]
    document.add_heading("先把现状说清楚", level=1)
    document.add_paragraph(display_text(situation.get("summary", "-")))
    add_docx_table(
        document,
        ["目标", "期限", "成功标准", "阶段"],
        [[situation.get("goal", "-"), situation.get("deadline", "-"), situation.get("success_metric", "-"), situation.get("stage", "-")]],
    )
    document.add_heading("现状够不够清楚", level=1)
    clarity = report["current_state_clarity"]
    document.add_paragraph(f"能不能开始判断：{'可以继续判断关键卡点' if clarity['diagnosis_allowed'] else '还不能直接下结论，先追问'}")
    document.add_paragraph(display_text(clarity["snapshot"]))
    add_docx_table(document, ["维度", "得分"], [[CLARITY_DIMENSION_LABELS.get(key, key), f"{value:.0f}"] for key, value in clarity["dimension_scores"].items()])
    document.add_heading("如果还不够清楚，先问这些", level=2)
    add_docx_bullets(document, clarity.get("questions") or ["现状已经够用，下一步只需要让用户确认上面的复述是否准确。"])
    document.add_heading("哪些是事实，哪些还只是判断", level=2)
    add_docx_table(document, ["事实", "类型", "备注"], [[fact["text"], source_label(fact["source_type"]), fact.get("note") or "-"] for fact in report["facts"]])
    analysis_logic = report["analysis_logic"]
    decision_model = report["decision_model"]
    document.add_heading("一张图看懂：从表象到主要矛盾", level=1)
    document.add_paragraph(display_text(report["visuals"]["summary"]))
    document.add_heading("分析流程图", level=2)
    document.add_paragraph("用户描述 -> 追问补全现状 -> 看得见的问题 -> 上升一层 -> 看不见的根部变量 -> 贝叶斯更新 -> 奥卡姆剃刀 -> 主要矛盾/次要矛盾 -> 行动与复盘")
    document.add_heading("冰山模型", level=2)
    add_docx_table(
        document,
        ["层级", "内容"],
        [
            ["水面上：看得见的问题", "；".join(analysis_logic["visible_problems"][:4])],
            ["水面下：看不见的根部变量", "；".join(analysis_logic["hidden_root_variables"][:4])],
            ["最底层：当前主要矛盾", report["principal"].get("name", "-")],
        ],
    )
    document.add_heading("贝叶斯更新与奥卡姆剃刀", level=2)
    document.add_paragraph(display_text(decision_model["bayesian_note"]))
    document.add_paragraph(display_text(decision_model["ockham_note"]))
    add_docx_table(
        document,
        ["候选矛盾", "更新后可信度", "解释力", "简洁度"],
        [[row["name"], pct(row.get("posterior")), pct(row.get("likelihood")), pct(row.get("parsimony"))] for row in decision_model["candidates"][:5]],
    )
    document.add_heading("主要矛盾判断过程", level=1)
    document.add_paragraph(display_text(analysis_logic["method_note"]))
    document.add_paragraph(f"第一性原理追问：{display_text(analysis_logic['first_principles_question'])}")
    document.add_heading("看得见的问题", level=2)
    add_docx_bullets(document, analysis_logic["visible_problems"])
    document.add_heading("上升一层：找看不见的根部变量", level=2)
    document.add_paragraph(display_text(analysis_logic["higher_level_logic"]))
    add_docx_bullets(document, analysis_logic["hidden_root_variables"])
    document.add_heading("为什么不是先处理表面问题", level=2)
    document.add_paragraph(display_text(analysis_logic["why_not_surface"]))
    document.add_heading("候选矛盾怎么排序", level=2)
    document.add_paragraph(display_text(analysis_logic["ranking_rule"]))
    add_docx_table(
        document,
        ["排序", "可能的卡点", "大致类型", "优先级", "依据"],
        [[idx + 1, row.get("name", "-"), plain_type_label(row.get("type")), f"{row['weighted_total']:.2f}/5", "；".join(row.get("evidence") or ["-"])] for idx, row in enumerate(report["contradictions"])],
    )
    principal = report["principal"]
    document.add_heading("主要矛盾（最关键的卡点）", level=1)
    document.add_paragraph(f"主要矛盾：{display_text(principal.get('name', '-'))}")
    document.add_paragraph(f"主要方面（现在最影响局面的一侧）：{display_text(principal.get('principal_aspect') or '需要继续追问')}")
    document.add_paragraph(f"为什么先处理它：{display_text(principal.get('why_principal') or '按现有事实看，它最影响目标推进。')}")
    document.add_paragraph(f"什么情况说明判断要改：{display_text(principal.get('reversal_condition') or '当关键指标改善但结果仍未变化时，需要重新判断。')}")
    document.add_heading("次要矛盾（先不主攻，但要盯住）", level=1)
    add_docx_table(document, ["次要矛盾", "为什么暂时不主攻", "什么时候要处理"], [[row.get("name", "-"), row.get("defer_reason") or "当前不是资源投放主线。", row.get("monitor_trigger") or "阶段变化时再看。"] for row in report["secondary"]])
    allocation = report["resource_allocation"]
    document.add_heading("时间、精力、资源应该怎么重新分配", level=1)
    document.add_paragraph(display_text(allocation.get("note", "-")))
    add_docx_table(
        document,
        ["资源项", "当前分配", "建议分配", "变化方向"],
        [
            [
                item.get("label", "-"),
                f"{as_float(item.get('current')):.0f}{allocation.get('unit', '%')}",
                f"{as_float(item.get('recommended')):.0f}{allocation.get('unit', '%')}",
                f"{as_float(item.get('recommended')) - as_float(item.get('current')):+.0f}{allocation.get('unit', '%')}",
            ]
            for item in allocation.get("items", [])
        ],
    )
    document.add_heading("接下来怎么做", level=1)
    add_docx_table(
        document,
        ["动作", "谁来做", "什么时候", "要用什么", "做到什么算完成", "做成把握", "预计帮助"],
        [[row.get("action", "-"), row.get("owner", "-"), row.get("deadline", "-"), row.get("resource", "-"), row.get("metric", "-"), pct(row["completion_probability"]), pct(row["expected_uplift"])] for row in report["actions"]],
    )
    projection = report["projection"]
    document.add_heading("做完以后可能怎样", level=1)
    add_docx_table(document, ["不专门处理时", "这些动作可能带来的帮助", "风险可能拖后腿", "按建议做后的估计"], [[pct(projection["baseline"]), pct(projection["expected_action_uplift"]), pct(projection["expected_risk_drag"]), pct(projection["recommended_probability"])]])
    add_docx_table(document, ["情况", "大概概率", "怎么理解"], [[item["name"], pct(item["probability"]), item["reading"]] for item in projection["scenarios"]])
    document.add_heading("哪些变化会影响这个估计", level=2)
    add_docx_bullets(document, projection.get("sensitivity_notes") or ["行动后概率依赖用户补充证据和执行反馈。"])
    stage_transition = report["stage_transition"]
    document.add_heading("主要矛盾什么时候会转移", level=1)
    document.add_paragraph(display_text(stage_transition.get("note", "-")))
    add_docx_table(document, ["阶段", "可能主导的矛盾", "触发条件"], [[item.get("stage", "-"), item.get("principal", "-"), item.get("trigger", "-")] for item in stage_transition.get("stages", [])])
    document.add_heading("什么时候回头看", level=1)
    document.add_paragraph(f"复盘时间：{report['review'].get('review_time', '7到30天后复盘')}")
    add_docx_bullets(document, report["review"].get("review_conditions") or [])
    document.add_heading("注意事项", level=1)
    add_docx_bullets(document, report["warnings"] or ["本报告只用于结构化诊断和行动辅助，不替代专业意见。"])
    document.save(path)


def write_docx_from_markdown(markdown_path: Path, docx_path: Path) -> None:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("python-docx is not installed and pandoc is not available")
    completed = subprocess.run([pandoc, str(markdown_path), "-o", str(docx_path)], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if completed.returncode != 0:
        raise RuntimeError(completed.stderr.strip() or "pandoc DOCX fallback failed")


def write_pdf(html_path: Path, pdf_path: Path) -> None:
    if HTML is not None:
        HTML(filename=str(html_path)).write_pdf(str(pdf_path))
        return
    python = shutil.which("python3")
    if python:
        code = (
            "from weasyprint import HTML; import sys; "
            "HTML(filename=sys.argv[1]).write_pdf(sys.argv[2])"
        )
        completed = subprocess.run([python, "-c", code, str(html_path), str(pdf_path)], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if completed.returncode == 0:
            return
        raise RuntimeError(completed.stderr.strip() or "system python3 could not render PDF with WeasyPrint")
    raise RuntimeError("WeasyPrint is not available")


def write_bundle(request_path: Path, output_dir: Path) -> dict[str, Any]:
    request = load_request(request_path)
    report = build_report(request)
    output_dir.mkdir(parents=True, exist_ok=True)
    slug = slugify(report["slug"])
    paths = {
        "json": output_dir / f"{slug}.report.json",
        "markdown": output_dir / f"{slug}.md",
        "html": output_dir / f"{slug}.html",
        "docx": output_dir / f"{slug}.docx",
        "pdf": output_dir / f"{slug}.pdf",
    }
    warnings: list[str] = []
    paths["json"].write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    paths["markdown"].write_text(render_markdown(report), encoding="utf-8")
    paths["html"].write_text(render_html(report), encoding="utf-8")
    try:
        write_docx(report, paths["docx"])
    except Exception as exc:
        try:
            write_docx_from_markdown(paths["markdown"], paths["docx"])
            warnings.append(f"DOCX generated through pandoc fallback because direct Word rendering failed: {exc}")
        except Exception as fallback_exc:
            warnings.append(f"DOCX generation failed: {exc}; fallback failed: {fallback_exc}")
    try:
        write_pdf(paths["html"], paths["pdf"])
    except Exception as exc:
        warnings.append(f"PDF generation failed: {exc}")
    report["generation"] = {
        "paths": {key: str(path) for key, path in paths.items() if path.exists()},
        "warnings": warnings,
    }
    paths["json"].write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report["generation"]


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a Yao Crux report bundle from JSON.")
    parser.add_argument("input_json", type=Path)
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()
    generation = write_bundle(args.input_json, args.output_dir)
    for kind, path in generation["paths"].items():
        print(f"{kind}: {path}")
    for warning in generation["warnings"]:
        print(f"warning: {warning}", file=sys.stderr)
    return 0 if "markdown" in generation["paths"] and "html" in generation["paths"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
